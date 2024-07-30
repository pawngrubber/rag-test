import pytest
from unittest.mock import MagicMock
from llm import VectorDB, read_txt, read_pdf, read_docx, process_and_add_files  # Replace 'my_module' with the actual module name

@pytest.fixture
def mock_vector_db():
    vector_db = VectorDB()
    vector_db.add_texts = MagicMock()
    vector_db.query_texts = MagicMock(return_value=["This is a test document."])
    return vector_db

def test_read_txt(tmp_path):
    # Create a temporary text file
    file_path = tmp_path / "test.txt"
    file_path.write_text("This is a test text file.")
    
    # Test read_txt function
    content = read_txt(file_path)
    assert content == "This is a test text file."

def test_read_pdf(tmp_path):
    # Create a temporary PDF file
    from reportlab.pdfgen import canvas
    file_path = tmp_path / "test.pdf"
    c = canvas.Canvas(str(file_path))
    c.drawString(100, 750, "This is a test PDF file.")
    c.save()
    
    # Test read_pdf function
    content = read_pdf(file_path)
    assert "This is a test PDF file." in content

def test_read_docx(tmp_path):
    # Create a temporary DOCX file
    from docx import Document
    file_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("This is a test DOCX file.")
    doc.save(file_path)
    
    # Test read_docx function
    content = read_docx(file_path)
    assert "This is a test DOCX file." in content

def test_process_and_add_files(mock_vector_db, tmp_path):
    # Create temporary files
    txt_path = tmp_path / "test.txt"
    txt_path.write_text("This is a test text file.")
    
    pdf_path = tmp_path / "test.pdf"
    from reportlab.pdfgen import canvas
    c = canvas.Canvas(str(pdf_path))
    c.drawString(100, 750, "This is a test PDF file.")
    c.save()
    
    docx_path = tmp_path / "test.docx"
    from docx import Document
    doc = Document()
    doc.add_paragraph("This is a test DOCX file.")
    doc.save(docx_path)
    
    file_paths = [str(txt_path), str(pdf_path), str(docx_path)]
    
    # Test process_and_add_files function
    process_and_add_files(mock_vector_db, file_paths)
    
    # Check if add_texts was called with the correct arguments
    mock_vector_db.add_texts.assert_called()
    args, _ = mock_vector_db.add_texts.call_args
    assert len(args[0]) == 3
    assert "This is a test text file." in args[0]
    assert "This is a test PDF file." in args[0]
    assert "This is a test DOCX file." in args[0]

def test_query_texts(mock_vector_db):
    query = "test"
    results = mock_vector_db.query_texts(query)
    assert results == ["This is a test document."]

if __name__ == "__main__":
    pytest.main()
